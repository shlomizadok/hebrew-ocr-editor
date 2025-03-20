from flask import Flask, request, jsonify, render_template, send_file
import os
import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document
import tempfile
import json
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Set Tesseract data path
os.environ['TESSDATA_PREFIX'] = '/opt/homebrew/share/tessdata/'

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


def process_image(image_path, area=None):
  try:
    # Read the image
    img = cv2.imread(image_path)
    if img is None:
      logger.error("Failed to read image")
      return []

    logger.debug(f"Image shape: {img.shape}")

    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    if area:
      # Extract the selected area
      x, y, w, h = map(int, [area['x'], area['y'], area['width'], area['height']])
      logger.debug(f"Processing area: x={x}, y={y}, w={w}, h={h}")

      # Ensure coordinates are within image bounds
      x = max(0, min(x, img.shape[1]))
      y = max(0, min(y, img.shape[0]))
      w = min(w, img.shape[1] - x)
      h = min(h, img.shape[0] - y)

      # Extract the region of interest
      roi = gray[y:y+h, x:x+w]

      # Save original ROI for debugging
      cv2.imwrite('debug_original_roi.png', roi)

      # 1. Resize if needed (scale up small text)
      scale = 1
      if w < 200 or h < 200:
        scale = 2
        roi = cv2.resize(roi, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

      # 2. Remove noise
      roi = cv2.fastNlMeansDenoising(roi)

      # 3. Increase contrast
      roi = cv2.convertScaleAbs(roi, alpha=1.5, beta=0)

      # 4. Apply adaptive thresholding
      roi = cv2.adaptiveThreshold(
          roi,
          255,
          cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
          cv2.THRESH_BINARY,
          21,  # Block size
          11   # C constant
      )

      # 5. Dilate to make text more prominent
      kernel = np.ones((2, 2), np.uint8)
      roi = cv2.dilate(roi, kernel, iterations=1)

      # Save processed ROI for debugging
      cv2.imwrite('debug_processed_roi.png', roi)
      logger.debug("Saved debug images")

      # Use Hebrew language for OCR with additional configuration
      custom_config = r'--oem 3 --psm 6 -l heb'
      text = pytesseract.image_to_string(roi, config=custom_config)
      logger.debug(f"OCR Result: {text}")

      return [{
          'x': area['x'],
          'y': area['y'],
          'width': area['width'],
          'height': area['height'],
          'text': text.strip()
      }]
    else:
      # Process entire image if no area specified
      text = pytesseract.image_to_string(gray, lang='heb')
      return [{
          'x': 0,
          'y': 0,
          'width': img.shape[1],
          'height': img.shape[0],
          'text': text.strip()
      }]
  except Exception as e:
    logger.error(f"Error in process_image: {str(e)}", exc_info=True)
    raise


@app.route('/')
def index():
  return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
  if 'file' not in request.files:
    return jsonify({'error': 'No file part'}), 400

  file = request.files['file']
  if file.filename == '':
    return jsonify({'error': 'No selected file'}), 400

  if file:
    filename = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filename)

    try:
      area = None
      if 'area' in request.form:
        area = json.loads(request.form['area'])
        logger.debug(f"Received area: {area}")

      areas = process_image(filename, area)
      return jsonify({'areas': areas})
    except Exception as e:
      logger.error(f"Error in upload_file: {str(e)}", exc_info=True)
      return jsonify({'error': str(e)}), 500
    finally:
      # Clean up the uploaded file
      os.remove(filename)


@app.route('/save-docx', methods=['POST'])
def save_docx():
  data = request.json
  areas = data.get('areas', [])
  original_filename = data.get('filename', 'output')

  try:
    # Create DOCX
    doc = Document()
    for area in areas:
      paragraph = doc.add_paragraph()
      # Set RTL alignment for Hebrew text
      paragraph.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
      paragraph.add_run(area['text'])

    # Save to temporary file
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_docx.name)

    # Send the DOCX file
    return send_file(
        temp_docx.name,
        as_attachment=True,
        download_name=f"{original_filename}.docx"
    )
  except Exception as e:
    logger.error(f"Error in save_docx: {str(e)}", exc_info=True)
    return jsonify({'error': str(e)}), 500
  finally:
    # Clean up temporary file
    try:
      os.unlink(temp_docx.name)
    except:
      pass


if __name__ == '__main__':
  app.run(debug=True)
